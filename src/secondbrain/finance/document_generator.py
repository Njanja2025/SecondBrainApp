"""
Document generation module for SecondBrain
"""

from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Union, Any
from decimal import Decimal
import logging
import json
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from fpdf import FPDF
import matplotlib.pyplot as plt
import seaborn as sns
from dataclasses import dataclass
from enum import Enum

# Configure logging
logger = logging.getLogger(__name__)


class ReportType(Enum):
    """Types of income reports."""

    DAILY = "daily"
    WEEKLY = "weekly"
    MONTHLY = "monthly"
    CUSTOM = "custom"


class DocumentStyle(Enum):
    """Document styling options."""

    MODERN = "modern"
    CLASSIC = "classic"
    MINIMAL = "minimal"
    CORPORATE = "corporate"


@dataclass
class DocumentConfig:
    """Configuration for document generation."""

    title: str
    style: DocumentStyle = DocumentStyle.MODERN
    include_charts: bool = True
    include_summary: bool = True
    include_notes: bool = True
    output_dir: Optional[str] = None
    template_path: Optional[str] = None
    watermark: Optional[str] = None
    company_logo: Optional[str] = None


class DocumentGenerator:
    """Base class for document generation."""

    def __init__(self, config: DocumentConfig):
        """
        Initialize the document generator.

        Args:
            config: Document configuration
        """
        self.config = config
        self.output_dir = (
            Path(config.output_dir)
            if config.output_dir
            else Path.home() / ".secondbrain" / "reports"
        )
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self._load_template()

    def _load_template(self):
        """Load document template if specified."""
        if self.config.template_path:
            try:
                with open(self.config.template_path, "r") as f:
                    self.template = json.load(f)
            except Exception as e:
                logger.warning(f"Failed to load template: {e}")
                self.template = {}
        else:
            self.template = {}

    def _format_currency(self, amount: Decimal) -> str:
        """Format currency amount with proper formatting."""
        return f"${amount:,.2f}"

    def _get_style_colors(self) -> Dict[str, RGBColor]:
        """Get color scheme based on document style."""
        colors = {
            DocumentStyle.MODERN: {
                "primary": RGBColor(41, 128, 185),  # Blue
                "secondary": RGBColor(52, 152, 219),  # Light Blue
                "accent": RGBColor(231, 76, 60),  # Red
            },
            DocumentStyle.CLASSIC: {
                "primary": RGBColor(44, 62, 80),  # Dark Blue
                "secondary": RGBColor(52, 73, 94),  # Gray
                "accent": RGBColor(192, 57, 43),  # Dark Red
            },
            DocumentStyle.MINIMAL: {
                "primary": RGBColor(52, 73, 94),  # Gray
                "secondary": RGBColor(149, 165, 166),  # Light Gray
                "accent": RGBColor(231, 76, 60),  # Red
            },
            DocumentStyle.CORPORATE: {
                "primary": RGBColor(41, 128, 185),  # Blue
                "secondary": RGBColor(52, 152, 219),  # Light Blue
                "accent": RGBColor(46, 204, 113),  # Green
            },
        }
        return colors.get(self.config.style, colors[DocumentStyle.MODERN])


class WordDocumentGenerator(DocumentGenerator):
    """Word document generator for income reports."""

    def generate_income_report(self, data: List[Dict[str, str]]) -> str:
        """
        Generate a Word document income report.

        Args:
            data: List of dictionaries containing income data

        Returns:
            Path to generated document
        """
        try:
            doc = Document()
            colors = self._get_style_colors()

            # Add company logo if specified
            if self.config.company_logo and os.path.exists(self.config.company_logo):
                doc.add_picture(self.config.company_logo, width=Inches(2))
                doc.add_paragraph()

            # Add title with styling
            title = doc.add_heading(self.config.title, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.style.font.color.rgb = colors["primary"]

            # Add date
            date_para = doc.add_paragraph(
                f"Date: {datetime.now().strftime('%Y-%m-%d')}"
            )
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_para.style.font.color.rgb = colors["secondary"]
            doc.add_paragraph()

            # Add income table
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Add header row with styling
            header_cells = table.rows[0].cells
            for cell in header_cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.paragraphs[0].style.font.color.rgb = colors["primary"]

            header_cells[0].text = "Income Source"
            header_cells[1].text = "Details"
            header_cells[2].text = "Amount (USD)"

            # Add data rows
            total = Decimal("0")
            for entry in data:
                revenue = Decimal(str(entry["Revenue"]).replace(",", ""))
                total += revenue

                row_cells = table.add_row().cells
                row_cells[0].text = entry["Source"]
                row_cells[1].text = entry["Item"]
                row_cells[2].text = self._format_currency(revenue)

                # Add notes if enabled
                if self.config.include_notes and "Notes" in entry:
                    note_cell = table.add_row().cells[0]
                    note_cell.merge(table.rows[-1].cells[1])
                    note_cell.merge(table.rows[-1].cells[2])
                    note_cell.text = f"Note: {entry['Notes']}"
                    note_cell.paragraphs[0].style.font.italic = True
                    note_cell.paragraphs[0].style.font.color.rgb = colors["secondary"]

            # Add total row with styling
            total_row = table.add_row().cells
            total_row[0].text = "Total Daily Revenue"
            total_row[1].text = ""
            total_row[2].text = self._format_currency(total)
            total_row[0].paragraphs[0].style.font.bold = True
            total_row[2].paragraphs[0].style.font.bold = True
            total_row[2].paragraphs[0].style.font.color.rgb = colors["accent"]

            # Add summary section if enabled
            if self.config.include_summary:
                doc.add_heading("Summary", level=1)
                summary = doc.add_paragraph()
                summary.add_run("Total Revenue: ").bold = True
                summary.add_run(self._format_currency(total))
                summary.add_run("\nNumber of Sources: ").bold = True
                summary.add_run(str(len(data)))
                summary.add_run("\nAverage Revenue per Source: ").bold = True
                summary.add_run(self._format_currency(total / len(data)))

            # Add charts if enabled
            if self.config.include_charts:
                self._add_charts(doc, data)

            # Add footer
            doc.add_paragraph()
            footer = doc.add_paragraph("Autogenerated by SecondBrainApp | Phantom AI")
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer.style.font.color.rgb = colors["secondary"]

            # Add watermark if specified
            if self.config.watermark:
                self._add_watermark(doc, self.config.watermark)

            # Save document
            output_path = (
                self.output_dir
                / f"income_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            )
            doc.save(str(output_path))

            logger.info(f"Word document generated: {output_path}")
            return str(output_path)

        except Exception as e:
            logger.error(f"Failed to generate Word document: {e}")
            raise

    def _add_charts(self, doc: Document, data: List[Dict[str, str]]):
        """Add charts to the document."""
        try:
            # Create revenue distribution pie chart
            plt.figure(figsize=(8, 6))
            revenues = [Decimal(str(d["Revenue"]).replace(",", "")) for d in data]
            sources = [d["Source"] for d in data]
            plt.pie(revenues, labels=sources, autopct="%1.1f%%")
            plt.title("Revenue Distribution by Source")

            # Save chart
            chart_path = self.output_dir / "temp_chart.png"
            plt.savefig(chart_path, dpi=300, bbox_inches="tight")
            plt.close()

            # Add chart to document
            doc.add_heading("Revenue Distribution", level=2)
            doc.add_picture(str(chart_path), width=Inches(6))

            # Clean up
            os.remove(chart_path)

        except Exception as e:
            logger.error(f"Failed to add charts: {e}")

    def _add_watermark(self, doc: Document, text: str):
        """Add watermark to the document."""
        try:
            section = doc.sections[0]
            footer = section.footer
            paragraph = footer.paragraphs[0]
            paragraph.text = text
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.style.font.color.rgb = RGBColor(200, 200, 200)

        except Exception as e:
            logger.error(f"Failed to add watermark: {e}")


class PDFDocumentGenerator(DocumentGenerator):
    """PDF document generator for income reports."""

    def generate_income_report(self, data: List[Dict[str, str]]) -> str:
        """
        Generate a PDF income report.

        Args:
            data: List of dictionaries containing income data

        Returns:
            Path to generated document
        """
        try:
            pdf = FPDF()
            pdf.add_page()

            # Add company logo if specified
            if self.config.company_logo and os.path.exists(self.config.company_logo):
                pdf.image(self.config.company_logo, x=10, y=10, w=40)
                pdf.ln(45)

            # Add header
            pdf.set_font("Arial", "B", 14)
            pdf.cell(0, 10, self.config.title, ln=True, align="C")
            pdf.set_font("Arial", "", 12)
            pdf.cell(
                0,
                10,
                f"Date: {datetime.now().strftime('%Y-%m-%d')}",
                ln=True,
                align="C",
            )
            pdf.ln(10)

            # Add income table
            pdf.set_font("Arial", "B", 12)
            pdf.cell(60, 10, "Income Source", border=1)
            pdf.cell(60, 10, "Details", border=1)
            pdf.cell(60, 10, "Amount (USD)", border=1)
            pdf.ln()

            # Add data rows
            pdf.set_font("Arial", "", 12)
            total = Decimal("0")
            for entry in data:
                revenue = Decimal(str(entry["Revenue"]).replace(",", ""))
                total += revenue

                pdf.cell(60, 10, entry["Source"], border=1)
                pdf.cell(60, 10, entry["Item"], border=1)
                pdf.cell(60, 10, self._format_currency(revenue), border=1)
                pdf.ln()

                # Add notes if enabled
                if self.config.include_notes and "Notes" in entry:
                    pdf.set_font("Arial", "I", 10)
                    pdf.cell(180, 10, f"Note: {entry['Notes']}", border=0)
                    pdf.ln()
                    pdf.set_font("Arial", "", 12)

            # Add total row
            pdf.ln(10)
            pdf.set_font("Arial", "B", 12)
            pdf.cell(120, 10, "Total Daily Revenue", border=1)
            pdf.cell(60, 10, self._format_currency(total), border=1)

            # Add summary section if enabled
            if self.config.include_summary:
                pdf.ln(20)
                pdf.set_font("Arial", "B", 12)
                pdf.cell(0, 10, "Summary", ln=True)
                pdf.set_font("Arial", "", 12)
                pdf.cell(
                    0, 10, f"Total Revenue: {self._format_currency(total)}", ln=True
                )
                pdf.cell(0, 10, f"Number of Sources: {len(data)}", ln=True)
                pdf.cell(
                    0,
                    10,
                    f"Average Revenue per Source: {self._format_currency(total / len(data))}",
                    ln=True,
                )

            # Add charts if enabled
            if self.config.include_charts:
                self._add_charts(pdf, data)

            # Add footer
            pdf.set_y(-15)
            pdf.set_font("Arial", "I", 8)
            pdf.cell(0, 10, "Autogenerated by SecondBrainApp | Phantom AI", align="C")

            # Add watermark if specified
            if self.config.watermark:
                self._add_watermark(pdf, self.config.watermark)

            # Save document
            output_path = (
                self.output_dir
                / f"income_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            )
            pdf.output(str(output_path))

            logger.info(f"PDF document generated: {output_path}")
            return str(output_path)

        except Exception as e:
            logger.error(f"Failed to generate PDF document: {e}")
            raise

    def _add_charts(self, pdf: FPDF, data: List[Dict[str, str]]):
        """Add charts to the PDF."""
        try:
            # Create revenue distribution pie chart
            plt.figure(figsize=(8, 6))
            revenues = [Decimal(str(d["Revenue"]).replace(",", "")) for d in data]
            sources = [d["Source"] for d in data]
            plt.pie(revenues, labels=sources, autopct="%1.1f%%")
            plt.title("Revenue Distribution by Source")

            # Save chart
            chart_path = self.output_dir / "temp_chart.png"
            plt.savefig(chart_path, dpi=300, bbox_inches="tight")
            plt.close()

            # Add chart to PDF
            pdf.ln(20)
            pdf.set_font("Arial", "B", 12)
            pdf.cell(0, 10, "Revenue Distribution", ln=True)
            pdf.image(str(chart_path), x=10, y=None, w=190)

            # Clean up
            os.remove(chart_path)

        except Exception as e:
            logger.error(f"Failed to add charts: {e}")

    def _add_watermark(self, pdf: FPDF, text: str):
        """Add watermark to the PDF."""
        try:
            pdf.set_font("Arial", "I", 60)
            pdf.set_text_color(200, 200, 200)
            pdf.set_xy(0, 0)
            pdf.cell(0, 297, text, 0, 0, "C")
            pdf.set_text_color(0, 0, 0)

        except Exception as e:
            logger.error(f"Failed to add watermark: {e}")


def generate_income_report(
    data: List[Dict[str, str]],
    config: Optional[DocumentConfig] = None,
    format: str = "pdf",
) -> str:
    """
    Generate an income report in the specified format.

    Args:
        data: List of dictionaries containing income data
        config: Document configuration (optional)
        format: Output format ("pdf" or "docx")

    Returns:
        Path to generated document
    """
    try:
        if config is None:
            config = DocumentConfig(
                title="SecondBrainApp Daily Income Summary", style=DocumentStyle.MODERN
            )

        if format.lower() == "pdf":
            generator = PDFDocumentGenerator(config)
        elif format.lower() == "docx":
            generator = WordDocumentGenerator(config)
        else:
            raise ValueError(f"Unsupported format: {format}")

        return generator.generate_income_report(data)

    except Exception as e:
        logger.error(f"Failed to generate income report: {e}")
        raise
